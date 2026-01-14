"""
文档处理器 - 多格式解析

支持的格式：
- 纯文本 (.txt)
- Markdown (.md)
- PDF (.pdf)
- Word (.docx)
- HTML (.html)
"""

from typing import Any, Dict, List, Optional
from dataclasses import dataclass
import os
import uuid


@dataclass
class Document:
    """文档数据结构"""
    id: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[str] = None


class DocumentProcessor:
    """文档处理器 - 多格式解析与分块"""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    async def load(self, file_path: str) -> Document:
        """
        加载文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的文档
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".txt":
            content = await self._load_text(file_path)
        elif ext == ".md":
            content = await self._load_markdown(file_path)
        elif ext == ".pdf":
            content = await self._load_pdf(file_path)
        elif ext == ".docx":
            content = await self._load_docx(file_path)
        elif ext in [".html", ".htm"]:
            content = await self._load_html(file_path)
        else:
            content = await self._load_text(file_path)
        
        return Document(
            id=str(uuid.uuid4()),
            content=content,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "extension": ext,
            }
        )
    
    async def _load_text(self, file_path: str) -> str:
        """加载纯文本"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    async def _load_markdown(self, file_path: str) -> str:
        """加载Markdown"""
        return await self._load_text(file_path)
    
    async def _load_pdf(self, file_path: str) -> str:
        """加载PDF"""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")
    
    async def _load_docx(self, file_path: str) -> str:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text_parts = [para.text for para in doc.paragraphs]
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
    
    async def _load_html(self, file_path: str) -> str:
        """加载HTML"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            return soup.get_text(separator="\n")
        except ImportError:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")
    
    def chunk(self, document: Document) -> List[Dict[str, Any]]:
        """
        文档分块
        
        Args:
            document: 文档对象
            
        Returns:
            分块列表
        """
        text = document.content
        chunks = []
        
        start = 0
        chunk_idx = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # 尝试在句子边界分割
            if end < len(text):
                # 向后查找句子结束符
                for sep in ["。", ".", "!", "?", "\n"]:
                    pos = text.rfind(sep, start, end)
                    if pos > start:
                        end = pos + 1
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    "id": f"{document.id}_chunk_{chunk_idx}",
                    "content": chunk_text,
                    "metadata": {
                        **document.metadata,
                        "chunk_index": chunk_idx,
                        "start_char": start,
                        "end_char": end,
                    }
                })
                chunk_idx += 1
            
            start = end - self.chunk_overlap
        
        document.chunks = [c["content"] for c in chunks]
        return chunks
    
    async def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理文档：加载并分块
        
        Args:
            file_path: 文件路径
            
        Returns:
            分块列表
        """
        document = await self.load(file_path)
        return self.chunk(document)
    
    async def process_directory(
        self,
        dir_path: str,
        extensions: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        """
        处理目录下的所有文档
        
        Args:
            dir_path: 目录路径
            extensions: 要处理的文件扩展名
            
        Returns:
            所有分块
        """
        if extensions is None:
            extensions = [".txt", ".md", ".pdf", ".docx", ".html"]
        
        all_chunks = []
        
        for root, _, files in os.walk(dir_path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in extensions:
                    file_path = os.path.join(root, file)
                    try:
                        chunks = await self.process(file_path)
                        all_chunks.extend(chunks)
                    except Exception as e:
                        print(f"[DocumentProcessor] 处理失败 {file_path}: {e}")
        
        return all_chunks
